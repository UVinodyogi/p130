#!/usr/bin/env python
# coding: utf-8

# In[ ]:

import warnings
warnings.filterwarnings('ignore')
import os
import docx
import PyPDF2
import pdfplumber
import pikepdf
from PIL import Image 
import io 
import nltk
nltk.download('stopwords')
nltk.download('punkt')
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize 
import re
import pandas as pd
import streamlit as st
import plotly_express as px
import matplotlib.pyplot as plt
from sklearn.preprocessing import LabelEncoder
lb=LabelEncoder()

# loading the trained model
#pickle_in = open('classifier.pkl', 'rb') 
#classifier = pickle.load(pickle_in)




# Define key terms dictionary for fixing Role Applied for 
terms = {'WorkDay ERP':['workday', 'workday consultant', 'workday hcm', 'eib', 'picof', 
                        'workday studio','nnbound/outbound integrations'],
         'Peoplesoft':['peoplesoft', 'pia','ccb','birt','peci','ccw','pum','people tools',
                        'peoplesoft implementation','peoplesoft components',
                        'peoplesoft dba','peoplesoft admin','peoplesoft admin/dba','peopleSoft fscm', 
                        'peopletoolsupgrade','peopletools upgrade','process scheduler servers',
                        'peoplesoft hrms','peopleSoft consultant','peopledoft cloud',
                        'PeopleSoft migrations','eoplesoft Testing Framework','pure internet architecture'],             
         'Database Developer':['sql','sql server', 'ms sql server','msbi', 'sql developer', 'ssis','ssrs',
                        'ssms','t-sql','tsql','Razorsql', 'razor sql','triggers','powerbi','power bi',
                        'oracle sql', 'pl/sql', 'pl\sql','oracle', 'oracle 11g','oledb','cte','ddl',
                        'dml','etl','mariadb','maria db'],
         'Java Developer':['reactjs', 'react js', 'react js developer', 'html', 
                        'css3','xml','javascript','html5','boostrap','jquery', 'redux','php', 'node js',
                        'nodejs','apache','netbeans','nestjs','nest js','react developer','react hooks',
                        'jenkins'],
         'Data Scientist':['Data Science','Python','SQL','Machine Learning',]}

# List of all key terms to indicate skillset. Include all the key words in the list 
allTerms = ['workday', 'hcm', 'eib', 'picof','workday hcm',
                        'workday studio','nnbound/outbound integrations',
                        'peoplesoft', 'pia','ccb','birt','peci','ccw','pum','people tools',
                        'peoplesoft implementation','peoplesoft components',
                        'peoplesoft dba','peoplesoft admin','peoplesoft admin/dba','peopleSoft fscm', 
                        'peopletoolsupgrade','peopletools upgrade','process scheduler servers',
                        'peoplesoft hrms','peopleSoft consultant','peopledoft cloud',
                        'PeopleSoft migrations','eoplesoft Testing Framework','pure internet architecture',
                        'sql','sql server', 'ms sql server','msbi', 'sql developer', 'ssis','ssrs',
                        'ssms','t-sql','tsql','Razorsql', 'razor sql','triggers','powerbi','power bi',
                        'oracle sql', 'pl/sql', 'pl\sql','oracle', 'oracle 11g','oledb','cte','ddl',
                        'dml','etl','mariadb','maria db','reactjs', 'react js','React JS','ReactJS', 'react js developer', 'html', 
                        'css3','xml','javascript','html5','boostrap','jquery', 'redux','php', 'node js',
                        'nodejs','apache','netbeans','nestjs','nest js','react developer','react hooks',
                        'jenkins','datascience','python','data science','ML','machine learning','Machine Learning',
                        'Machine learning','AI','ai','Artificial Inteligence','artifitial ineligence','Modelling',
                        'Big Data','Bigdata','BigData','bigdata','AWS','Cloud Environments','oracle','Oracle']

# Function to extract text from resume
def getText(filename):
      
    # Create empty string 
    fullText = ''
    if filename.endswith('.docx'):
        doc = docx.Document(filename)
        
        for para in doc.paragraphs:
            fullText = fullText + para.text
            
           
    elif filename.endswith('.pdf'):
        with open(filename, "rb") as pdf_file:
            pdoc = PyPDF2.PdfFileReader(filename)
            number_of_pages = pdoc.getNumPages()
            page = pdoc.pages[0]
            page_content = page.extractText()
             
        for paragraph in page_content:
            fullText =  fullText + paragraph
        
            
    else:
        import aspose.words as aw
        output = aw.Document()
        # Remove all content from the destination document before appending.
        output.remove_all_children()
        input = aw.Document(filename)
        # Append the source document to the end of the destination document.
        output.append_document(input, aw.ImportFormatMode.KEEP_SOURCE_FORMATTING)
        output.save("Output.docx");
        doc = docx.Document('Output.docx')
        
        for para in doc.paragraphs:
            fullText = fullText + para.text
        fullText = fullText[79:]
         
    return (fullText)

# Function to remove punctuation and tokenize the text
def tokenText(extText):
   
    # Remove punctuation marks
    punc = '''!()-[]{};:'"\,.<>/?@#$%^&*_~'''
    for ele in extText:
        if ele in punc:
            puncText = extText.replace(ele, "")
            
    # Tokenize the text and remove stop words
    stop_words = set(stopwords.words('english'))
    puncText.split()
    word_tokens = word_tokenize(puncText)
    TokenizedText = [w for w in word_tokens if not w.lower() in stop_words]
    TokenizedText = []
  
    for w in word_tokens:
        if w not in stop_words:
            TokenizedText.append(w)
    return(TokenizedText)

# Function to read the tokenized text and search for the key words to dermine the Role Applied for
def roleApplied (Text):
    
    # covert the text to lower case
    for i in range(len(Text)):
        Text[i] = Text[i].lower()
    
    # Obtain the scores for each area
    for area in terms.keys():
        if area == 'WorkDay ERP':
            for word in terms[area]:
                if word in Text:
                    role = area
                    return (role)
                
        elif area == 'Peoplesoft':
            for word in terms[area]:
                if word in Text:
                    role = area
                    return(role)   
                
        elif area == 'Database Developer':
            for word in terms[area]:
                if word in Text:
                    role =  area
                    return(role)
            
        elif area == 'Java Developer':
             for word in terms[area]:
                if word in Text:
                    role = area
                    return(role)
        else:
            role = "Fresher"
            return(role)
        
# Function to extract Name and contact details
def contactDetails(Text):
    name = ''  
    for i in range(0,3):
        name = " ".join([name, Text[i]])
    return(name)

# Function to extract experience details
def expDetails(Text):
    global sent
   
    Text = Text.split()
   
    for i in range(len(Text)-2):
        Text[i].lower()
        
        if Text[i] ==  'years':
            sent =  Text[i-2] + ' ' + Text[i-1] +' ' + Text[i] +' '+ Text[i+1] +' ' + Text[i+2]
            return (sent)
        
        
# Function to extract skill set details
def skillSet(Text):
    t = []
    for i in range(len(Text)):
        if Text[i] in allTerms:
            if Text[i] in t:
                continue
            t.append(Text[i]) 
    return(t)
#to get phone numbers
def extract_mobile_number(text):
    phone = re.findall(re.compile(r'(?:(?:\+?([1-9]|[0-9][0-9]|[0-9][0-9][0-9])\s*(?:[.-]\s*)?)?(?:\(\s*([2-9]1[02-9]|[2-9][02-8]1|[2-9][02-8][02-9])\s*\)|([0-9][1-9]|[0-9]1[02-9]|[2-9][02-8]1|[2-9][02-8][02-9]))\s*(?:[.-]\s*)?)?([2-9]1[02-9]|[2-9][02-9]1|[2-9][02-9]{2})\s*(?:[.-]\s*)?([0-9]{4})(?:\s*(?:#|x\.?|ext\.?|extension)\s*(\d+))?'), text)
    
    if phone:
        number = ''.join(phone[0])
        if len(number) > 10:
            return '+' + number
        else:
            return number
#to get E-Mail 
def extract_email(email):
    email = re.findall("([^@|\s]+@[^@]+\.[^@|\s]+)", email)
    if email:
        try:
            return email[0].split()[0].strip(';')
        except IndexError:
            return None

def remove_dups_words(row):
    sentences = set(row.split(","))
    new_str = ','.join(sentences)
    return new_str

# this is the main function in which we define our webpage  
def jls_extract_def():
    
    return 


def main():
    st.set_page_config(
        page_title= "RESUME CLASSIFICATION",
        layout= "wide"
    )

    st.markdown("""
    # RESUME CLASSIFICATION
    
    This Simple Resume Analyzer was developed to analyze the resumes and to select the desired candidate/candidates 
    from a bunch of resumes by passing the folder path bellow and puch on the bellow buttons to get result.""")

    # following lines create boxes in which user can enter data required to make prediction 
    with st.form(key="form1"):
        st.warning(body="Supported file Formats: 'docx'or'doc'or'pdf' or all together")
        name=st.text_input(label= "Enter the folder path given below in the box")
        st.markdown("C:\\Users\\Admin\\Downloads\\Resumes ")
        submit = st.form_submit_button(label="SUBMIT")
        st.markdown("""Note: ignore this WinError 3""")

    path = name
    #if path is not None:
    #    try:
     #       st.write("Hey Buddy please provide appopriate path")
      #  except :
       #     pass
            #st.write("Hey Buddy please provide appopriate path")
    #path = st.text_input('Enter the resumes folder path')

    # Create an empt Data Frame ResumeText with two columns
    ResumeText = pd.DataFrame([], columns=['Name', 'Experience', 'skills','RoleApplied','Phone',"E-Mail"])
    
    # when 'Predict' is clicked, make the prediction and store it 
    #if st.button("Upload and Get Result"): 
        # Search the directory path and loop through the resume documents and callthe function getText
    #if st.button("Process"):
    for filename in os.listdir(path):
        filename = os.path.join(path, filename)
        extText = getText(filename)
        tokText = tokenText(extText)
        role = roleApplied(tokText)
        Name = contactDetails(tokText)
        experience = expDetails(extText)
        skills = skillSet(tokText)
        Phone = extract_mobile_number(extText)
        Email = extract_email(extText)
        NewRow = [Name,experience, skills,role,Phone,Email]
        ResumeText.loc[len(ResumeText)] = NewRow
        #st.dataframe(ResumeText)
        java = (ResumeText["RoleApplied"] == "Java Developer")
        #javares = ResumeText[java]
        workday = (ResumeText["RoleApplied"] == "WorkDay ERP")
        peosoft = (ResumeText["RoleApplied"] == "Peoplesoft")
        dbms = (ResumeText["RoleApplied"] == "Database Developer")
       
#    col1, col2, col3, col4, col5 = st.columns(5)

#    with col1:
#        if st.button("All Resumes"):
#            st.dataframe(ResumeText) 
#    with col2:
#        if st.button("JAVA Resumes"):
#            st.dataframe(ResumeText[java])              
#    with col3:
#        if st.button("DBMS Resumes"):
#            st.dataframe(ResumeText[dbms])
#    with col4:
#        if st.button("Peoplesoft Resumes"):
#            st.dataframe(ResumeText[peosoft])
#    with col5:
#        if st.button("Workday Resumes"):
#            st.dataframe(ResumeText[workday])
    
    df= ResumeText
    df['skills'] = df['skills'].apply(lambda x: ','.join(map(str, x)))
    df['skills'] = df["skills"].apply(remove_dups_words)
    df['Skills'] = df["skills"].value_counts()
    df['Skills'] = lb.fit_transform(df['skills'])
    df['role'] = lb.fit_transform(df["RoleApplied"])
    df["experience"]= df.Experience.str.extract("(\d.\d+|\d+)")

    #pie chart of all the skills.
    load= st.button("Load Data")
    #Initialize session state
    if "load_state" not in st.session_state :
        st.session_state.load_state = False

    if load or st.session_state.load_state:
        st.session_state.load_state= True
        st.write(df)
        
        st.subheader(body= "Have a look over the visualizations of skills and Role Applied")
        #user option
        opt = st.radio("Plot type :",["Pie chart","Scatter Plot & Heatmap","Histogram"])
        if opt == "Pie chart":
            col1,col2=st.columns(2)
            with col1:
                fig = px.pie(df,values="Skills",hole=0.0,labels={"label":"Skills","Skills":"Skills"},
                                hover_data=["skills"])#,custom_data=["RoleApplied"])
                fig.update_layout(title="Pie chart of Skills")
                fig.update_traces(textposition="inside",textinfo="percent+label")
                st.plotly_chart(fig)
            with col2:
                fig = px.pie(df,values="role",hole=0.0,labels={"label":"role","role":"role"},
                                hover_data=["Skills"],custom_data=["Skills","RoleApplied"])
                fig.update_layout(title="Pie chart of Job Roles")
                fig.update_traces(textposition= "inside", textinfo="label+percent",)
                                  #hovertemplate="RoleApplied:%{label}")
                st.plotly_chart(fig)
        elif opt == "Scatter Plot & Heatmap":
            col3, col4 = st.columns(2)
            with col3:
                fig = px.scatter(df,y= "skills", x="RoleApplied",title="Scatter Plot")
                st.plotly_chart(fig)
            with col4:
                fig = px.density_heatmap(df, x= "RoleApplied",y="skills",title="Heatmap")
                st.plotly_chart(fig)
        else :
            col5,col6=st.columns(2)
            with col5:
                fig = px.histogram(df,y="skills", x="role")
                fig.update_layout(title= "Sum of Role Applied regaring Skills")
                st.plotly_chart(fig)
            with col6:
                fig= px.histogram(df,x="RoleApplied", y="Skills")
                fig.update_layout(title = "Sum of Skills regarding Role Applied")
                st.plotly_chart(fig)


    lod = st.button(" Visualization on Experience , Skills and Roles ")
    choice = lod
    if choice == lod:
        if "lod_state" not in st.session_state:
            st.session_state.lod_state = False

        if lod or st.session_state.lod_state:
            st.session_state.lod_state= True
            st.subheader(body="Take a look at years of Experience & Skills")
            opt= st.radio("Choose plot type:",["Bar chart & Histogram","Heatmap & Scatterplot"])
            if opt== "Bar chart & Histogram":
                col7,col8=st.columns(2)
                with col7:
                    fig = px.bar(df,y="experience",x="RoleApplied",hover_data=["Experience"],
                         title="Bar chart of Experience")
                    st.plotly_chart(fig)
                with col8:
                    fig= px.histogram(df,x="experience",y="Skills",
                    title="Histogram of Experience(years) regaring skills")
                    st.plotly_chart(fig)
            else :
                col9,col10=st.columns(2)
                with col9:
                    fig= px.density_heatmap(df,x ="experience", y="skills",z="role",
                    title="Heatmap of Experience(years) regarding skills")
                    st.plotly_chart(fig)
                with col10:
                    fig = px.scatter(df,x="experience",y="skills",
                    title="Scatterplot of Experience(years) regarding skills")
                    st.plotly_chart(fig)
    

    daol = st.button("Find Desired Resume ")
    if "daol_state" not in st.session_state:
        st.session_state.daol_state = False

    if daol or st.session_state.daol_state:
        st.session_state.daol_state= True
        opt = st.radio("Choose candidate with prospective of :",["Skills","Experience(years)"])
        if opt == "Skills":
            Skill_option = df["skills"].unique().tolist()
            Skill = st.selectbox("Choose the candidate by selecting skills",Skill_option, 0)
        #Role_option = df["RoleApplied"].unique().tolist()
        #Role = st.selectbox("Select Role ",Role_option,0)
            df = df[df["skills"]==Skill] 
        #df2 = df[df["RoleApplied"]==Role]
            st.write(df)
            
        #st.write(df2)
        #    fig = px.pie(df,values="Skills",hole=0.3,labels={"label":"Skills","Skills":"skills"},
        #                    hover_data=["Experience"])
        #    fig.update_layout(title="Pie chart of Skills")
        #    fig.update_traces(textposition="inside",textinfo="percent+label")
        #    st.write(fig)
        else:
            #df2=df.sort_values(by="experience")
            experience_option = df["experience"].unique().tolist()
            exp = st.selectbox("Choose the candidate by selecting experience(years) ",experience_option,0)
            df = df[df['experience']==exp]
            st.write(df)

        col11,col12 =st.columns(2)
        with col11:
            fig= px.density_heatmap(df,x ="experience", y="skills",z="role")
            fig.update_layout(title="Heatmap for Experience(years) & skills of selected candidate")
            st.plotly_chart(fig)
        with col12:
            fig= px.scatter(df,x="experience",y="skills")
            fig.update_layout(title="Scatterplot for Experience(years) & skills of selected candidate")
            st.plotly_chart(fig)
        



if __name__=='__main__': 
    main()

